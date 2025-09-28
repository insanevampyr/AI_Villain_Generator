# duel_engine.py
# ------------------------------------------------------------
# Villain-vs-Villain Duel Engine — Cinematic Narration Mode
# Upgrades included here:
# • Violence hard-locked to max (2): extreme brutality & heavy swearing (no sexual content/hate slurs).
# • OTF: Over-the-Top Finishers use surviving props/terrain + decisive, gory finale with catchphrase.
# • CIC: Crippling Injury Continuity — auto-parsed injuries impose small scoring penalties & turn tax.
# • Dirty tricks + power variety bias each round.
# • Prop planning and per-round prop budget (3–4 foreground items).
# • Alias-first everywhere; continuity respected.
# • Auto-export each duel as a .docx to a local folder (see DUEL_DOCX_DIR).
# ------------------------------------------------------------

from __future__ import annotations
from dataclasses import dataclass, field
from typing import List, Dict, Optional, Tuple
import os, sys, json, random, textwrap, time, re, datetime as _dt

import argparse
from dotenv import load_dotenv

load_dotenv()

# --------------- Config knobs ---------------
ROUNDS = int(os.getenv("DUEL_ROUNDS", "10"))
USE_API = os.getenv("DUEL_USE_OPENAI", "1").strip().lower() in ("1","true","on")
OPENAI_MODEL = os.getenv("OPENAI_MODEL_NAME", "gpt-4o-mini")
TEMP = float(os.getenv("DUEL_TEMPERATURE", "0.85"))
VIOLENCE_LEVEL = 2  # hard-locked to max gore
WRAP = 92

PRICE_IN_PER_1K = float(os.getenv("PRICE_IN_PER_1K", "0.003"))
PRICE_OUT_PER_1K = float(os.getenv("PRICE_OUT_PER_1K", "0.006"))

# DOCX export directory (Windows path by default, configurable via env)
DUEL_DOCX_DIR = os.getenv(
    "DUEL_DOCX_DIR",
    r"C:\Users\VampyrLee\Desktop\AI_Villain\Duel Stories"
)

# --------------- Pronouns ---------------
PRONOUN_MAP = {
    "male":  ("He","him","his"),
    "female":("She","her","her"),
    "nonbinary":("They","them","their"),
    "other": ("They","them","their"),
}
def _pronouns_from_gender(gender: str) -> tuple[str,str,str]:
    return PRONOUN_MAP.get((gender or "").strip().lower(), PRONOUN_MAP["other"])

# --------------- Lair tag helpers ---------------
KNOWN_LAIR_TAGS = {
    "BioLab Sanctum": ["lab","biotech","low_light","glass","props","sterile"],
    "Acid Den Hideout": ["industrial","metal","steam","enclosed","props","toxic"],
}
AUTO_TAG_RULES = [
    ("lab", ["lab","biotech","glass","props","low_light"]),
    ("bio", ["lab","biotech","glass","props"]),
    ("acid", ["industrial","toxic","steam","metal","enclosed","props"]),
    ("den", ["industrial","metal","enclosed"]),
    ("factory", ["industrial","metal","props"]),
    ("catwalk", ["metal","height","enclosed"]),
    ("sewer", ["wet","enclosed","toxic"]),
]
def _auto_tags(lair_name: str) -> list[str]:
    n = (lair_name or "").lower()
    tags = set()
    for needle, tg in AUTO_TAG_RULES:
        if needle in n:
            tags.update(tg)
    if not tags:
        tags.update(["enclosed","props","low_light"])
    return sorted(tags)

# --------------- Data ---------------
@dataclass
class Villain:
    name: str
    alias: Optional[str] = None
    powers: List[str] = field(default_factory=list)
    weaknesses: List[str] = field(default_factory=list)
    catchphrase: Optional[str] = None
    vibe: Optional[str] = None
    subj: str = "They"
    obj: str = "them"
    pos: str = "their"
    def label(self) -> str:
        return self.alias or self.name

@dataclass
class Arena:
    name: str
    tags: List[str] = field(default_factory=list)

@dataclass
class RoundResult:
    r: int
    a_text: str
    b_text: str
    camera: Optional[str]
    a_delta: int
    b_delta: int
    a_total: int
    b_total: int

@dataclass
class DuelResult:
    a: Villain
    b: Villain
    arena: Arena
    scene_setter: str
    rounds: List[RoundResult]
    a_total: int
    b_total: int
    winner_label: str
    finisher: str

# --------------- Small utils ---------------
def _wrap(s: str) -> str:
    return textwrap.fill(s.strip(), width=WRAP)

def _bounded(n: int, lo: int, hi: int) -> int:
    return max(lo, min(hi, n))

def _name(v: Villain) -> str:
    return v.label()

# --------------- Momentum math ---------------
def _score_component(rng: random.Random, lo: int, hi: int) -> int:
    return rng.randint(lo, hi)

def _trait_bonus(actor: Villain, opp: Villain, arena: Arena) -> Dict[str, int]:
    p_text = " ".join(actor.powers).lower()
    opp_weak = " ".join(opp.weaknesses).lower()
    bonus = dict(A=0, C=0, S=0, R=0, E=0, D=0)
    if "toxin" in p_text or "acid" in p_text:
        if any(t in arena.tags for t in ("enclosed","metal","steam")):
            bonus["E"] += 1; bonus["S"] += 1
    if "shape" in p_text or "mimic" in p_text:
        if any(t in arena.tags for t in ("props","crowd","stage","lab","biotech")):
            bonus["S"] += 1
    if "freez" in opp_weak and ("ice" in p_text or "cold" in p_text):
        bonus["C"] += 2
    if "antitoxin" in opp_weak and ("toxin" in p_text or "acid" in p_text):
        bonus["C"] += 1
    return bonus

# --- CIC penalties based on injuries (lightweight, additive to stagger_penalty) ---
def _injury_penalty(label: str, ledger: dict, tactic_hint: str) -> int:
    injuries = (ledger.get("injuries", {}) or {}).get(label, []) or []
    pen = 0
    inj_txt = " ".join(injuries).lower()
    # Breath/ribs -> penalty on explosive strikes/mobility
    if "rib" in inj_txt or "breath" in inj_txt:
        pen += 1
    # Leg/limp (detected via hazards/openings wording as heuristic)
    limp_flags = ledger.get("openings", {}).get(label, "") or ""
    if any(k in limp_flags.lower() for k in ["stumble","off-balance","limp","knees","ankle","knee"]):
        pen += 1
    # Bleeding or deep lacerations -> cumulative fatigue
    if "bleeding" in inj_txt or "laceration" in inj_txt:
        pen += 1
    # Head trauma -> add another if tactic is mobility/grapple (coordination)
    if "concussion" in inj_txt and any(t in tactic_hint for t in ["mobility","grapple","feint"]):
        pen += 1
    return _bounded(pen, 0, 3)

def _round_delta(attacker: Villain, defender: Villain, arena: Arena,
                 rng: random.Random, stagger_penalty: int, combo: int,
                 extra_penalty: int = 0) -> Tuple[int, Dict[str,int]]:
    base = dict(
        A=_score_component(rng, 1, 4),
        C=_score_component(rng, -1, 3),
        S=_score_component(rng, 0, 2),
        R=_score_component(rng, -1, 2),
        E=_score_component(rng, 0, 2),
        D=_score_component(rng, 0, 3),
    )
    t = _trait_bonus(attacker, defender, arena)
    for k,v in t.items():
        base[k] = _bounded(base[k] + v, -3, 5)
    combo_bonus = min(combo, 3)
    penalty = stagger_penalty + extra_penalty
    delta = base["A"] + base["C"] + base["S"] + base["R"] + base["E"] - base["D"] + combo_bonus - penalty
    delta = _bounded(delta, -2, 12)
    base["combo"] = combo_bonus
    base["stagger"] = stagger_penalty
    base["injury_penalty"] = extra_penalty
    base["delta"] = delta
    return delta, base

# --------------- OpenAI IO + costs ---------------
COST_LOG = {"calls": [], "total_input_tokens": 0, "total_output_tokens": 0}
def _client():
    if not USE_API:
        print("[DEBUG] DUEL_USE_OPENAI not enabled.", file=sys.stderr)
        sys.exit(1)
    key = os.getenv("OPENAI_API_KEY","").strip()
    if not key:
        print("[DEBUG] OPENAI_API_KEY not set.", file=sys.stderr)
        sys.exit(1)
    try:
        from openai import OpenAI
        return OpenAI(api_key=key)
    except Exception as e:
        print(f"[DEBUG] OpenAI import/init failed: {e}", file=sys.stderr)
        sys.exit(1)

def _retry_call(fn, *args, **kwargs):
    attempts = 0; last_err = None
    while attempts < 3:
        try: return fn(*args, **kwargs)
        except Exception as e:
            last_err = e; attempts += 1; 
            if attempts < 3: time.sleep(0.7 * attempts)
    raise RuntimeError(str(last_err) if last_err else "Unknown OpenAI error")

def _chat_call(client, messages, max_tokens=500) -> str:
    def _do():
        resp = client.chat.completions.create(
            model=OPENAI_MODEL, messages=messages, temperature=TEMP, max_tokens=max_tokens
        )
        try:
            u = getattr(resp, "usage", None)
            in_tok  = int(getattr(u, "prompt_tokens", 0) or 0)
            out_tok = int(getattr(u, "completion_tokens", 0) or 0)
            COST_LOG["total_input_tokens"]  += in_tok
            COST_LOG["total_output_tokens"] += out_tok
            cost = (in_tok/1000.0)*PRICE_IN_PER_1K + (out_tok/1000.0)*PRICE_OUT_PER_1K
            COST_LOG["calls"].append({"in_tokens": in_tok, "out_tokens": out_tok, "cost_usd": round(cost, 6)})
        except Exception:
            pass
        return (resp.choices[0].message.content or "").strip()
    return _retry_call(_do)

# --------------- System prompts ---------------
VIOLENCE_TEXT = {
    0: "Keep violence gritty but not graphic; no gore. Mild swearing allowed.",
    1: "R-rated brutality allowed: blood spray, fractures, broken teeth, choking, concussions. Strong swearing allowed. No sexual content, no hate slurs.",
    2: "Extreme brutality allowed: shattered bones, pulped flesh, impalements, arterial spray, dismemberment when plausible. Heavy swearing ok. No sexual content and no hate slurs under any circumstance."
}[ _bounded(VIOLENCE_LEVEL,0,2) ]

ROUND_SYSTEM = f"""You are narrating one round of a cinematic comic-book style duel.
Rules:
- Present tense, 2–3 vivid sentences per villain. {VIOLENCE_TEXT}
- Villains play by no rules: cheap shots, feints, traps, grapples, improvised weapons, environment abuse.
- Powers must be used in varied, inventive ways: offense, defense, mobility, area denial, grapples, counters, feints, amplifiers (ricochets, reflections), combos with props or terrain.
- Each villain must ATTACK or COUNTER the other; no posing.
- CONSEQUENCES: describe injuries (cuts, breaks, blood), positioning, and openings caused. Injuries persist.
- Thoughts only when they add strategy, pain, or resolve (in *italics*).
- One CAMERA cue at most; format exactly: CAMERA: <description>.
Continuity:
- You will receive a ledger of injuries, hazards, props, and positional disadvantages; escalate or respect them.
- Props are PHYSICAL: if used, show them being picked up/dragged/kicked/thrown; nothing pops into a hand.
- Positional disadvantage matters: prone/flanked/struck-from-behind requires visible recovery and can cost that fighter their turn.
- Per-round prop budget: 3–4 foreground items max; at most one prop per fighter unless already in-hand.
- Use only the props listed; do not invent new items.
Momentum mapping:
  • +2..+5 = graze/small positional win.
  • +6..+9 = clear damaging hit; blood drawn; stagger.
  • +10..+12 = brutal injury or big prop/environment payoff.
Output format (no extra text):
A: <2–3 sentences for the first villain named in input>
B: <2–3 sentences for the second villain named in input>
Optional one line: CAMERA: <short>.
"""

SCENE_SYSTEM = """You are writing the cinematic scene-setter for a comic duel.
- 3–5 sentences, present tense.
- Describe atmosphere, lighting, hazards, and plausible props/hazards.
- Optional single CAMERA line.
- No villain actions or catchphrases yet.
"""

FINISHER_SYSTEM = f"""Narrate the finishing move in 3–5 sentences, present tense.
- Winner unleashes a decisive, over-the-top brutal action (allowed level: {VIOLENCE_TEXT}).
- Peak the winner's power; optionally combine it with 1–2 surviving props or a terrain beat for amplification.
- Respect continuity: injuries, hazards, and any prop already in-hand should shape the method of destruction.
- The loser ends in a clear, final state (KO, maimed beyond response, or explicit kill per comic-book tone).
- End with the winner's catchphrase in *italics*.
- No sexual content; no hate slurs.
"""

# --------------- Message builders ---------------
def _scene_setter_messages(a, b, arena):
    return [
        {"role":"system","content":SCENE_SYSTEM},
        {"role":"user","content":f"""Villains:
- {a.label()} (powers: {', '.join(a.powers)})
- {b.label()} (powers: {', '.join(b.powers)})

Arena: {arena.name} (tags: {', '.join(arena.tags)})
"""}
    ]

def _round_user_prompt(a, b, arena, ledger, round_idx, tactic_hint_a:str, tactic_hint_b:str) -> str:
    a_label, b_label = a.label(), b.label()
    inj_a = ledger.get("injuries", {}).get(a_label, []) or []
    inj_b = ledger.get("injuries", {}).get(b_label, []) or []
    hazards = ledger.get("hazards", []) or []
    props = ledger.get("props_in_play", []) or []
    openings_a = (ledger.get("openings", {}).get(a_label) or "none").strip()
    openings_b = (ledger.get("openings", {}).get(b_label) or "none").strip()
    pos_disadv = ledger.get("positional_disadvantage", "") or "none"

    # Crippling notes for model awareness
    crip_a = " | Crippling effects: " + (", ".join(inj_a) if inj_a else "none")
    crip_b = " | Crippling effects: " + (", ".join(inj_b) if inj_b else "none")

    continuity_text = "\n".join([
        f"- Injuries — {a_label}: {', '.join(inj_a) if inj_a else 'none recorded'}",
        f"- Injuries — {b_label}: {', '.join(inj_b) if inj_b else 'none recorded'}",
        "- Hazards: " + (", ".join(hazards) if hazards else "none recorded"),
        "- Props in play: " + (", ".join(props) if props else "none recorded"),
        f"- Openings — {a_label}: {openings_a}",
        f"- Openings — {b_label}: {openings_b}",
        f"- Positional disadvantage: {pos_disadv}",
    ])

    sig = ledger.get("ARENA_SIGNATURE") or []
    sig_line = ("; ".join(sig[:5])) if sig else (", ".join(arena.tags) if arena.tags else "—")

    return (
        f"Round {round_idx}\n"
        f"ARENA: {arena.name} | Signature: {sig_line}\n\n"
        f"VILLAIN A (first block): {a_label} | pronouns {a.subj}/{a.obj}/{a.pos} | powers: {', '.join(a.powers)}{crip_a}\n"
        f"Tactic bias for {a_label}: {tactic_hint_a}\n\n"
        f"VILLAIN B (second block): {b_label} | pronouns {b.subj}/{b.obj}/{b.pos} | powers: {', '.join(b.powers)}{crip_b}\n"
        f"Tactic bias for {b_label}: {tactic_hint_b}\n\n"
        f"PROPS IN PLAY (3–4 max):\n" + (("- " + "\n- ".join(props)) if props else "none") + "\n\n" +
        "CONTINUITY (persist & escalate):\n" + continuity_text + "\n\n" +
        "OUTPUT EXACTLY:\nA: ...\nB: ...\nOptional: CAMERA: ..."
    )

def _round_messages(a, b, arena, ledger, round_idx, a_total, b_total, total_rounds, tactic_a, tactic_b):
    rounds_left = total_rounds - round_idx
    if a_total > b_total:
        tilt = f"{a.label()} leads by {a_total-b_total}. Trailer should chase swings; leader can deny/counter."
    elif b_total > a_total:
        tilt = f"{b.label()} leads by {b_total-a_total}. Trailer should chase swings; leader can deny/counter."
    else:
        tilt = "Scores tied. Press for a swing while guarding vs counters."

    scoring = f"SCORING_CONTEXT: Round {round_idx}/{total_rounds} (left: {rounds_left}). Totals — {a.label()}: {a_total}, {b.label()}: {b_total}. {tilt}"
    sys = ROUND_SYSTEM + "\n" + scoring
    user = _round_user_prompt(a, b, arena, ledger, round_idx, tactic_a, tactic_b)
    return [ {"role":"system","content":sys}, {"role":"user","content":user} ]

def _finisher_messages(winner, loser, arena, ledger):
    user = {
        "winner": {"label": winner.label(), "pronouns":[winner.subj,winner.obj,winner.pos], "powers": winner.powers, "catchphrase": winner.catchphrase},
        "loser":  {"label": loser.label(),  "pronouns":[loser.subj,loser.obj,loser.pos],   "powers": loser.powers},
        "arena": {"name": arena.name, "tags": arena.tags},
        "ledger": ledger
    }
    return [{"role":"system","content":FINISHER_SYSTEM},
            {"role":"user","content":json.dumps(user, ensure_ascii=False)}]

# --------------- Prop planning ---------------
def _prop_pack_messages(a, b, arena):
    sys = "You are a fight-scene planner. Return a compact plan; no narration."
    user = f"""Arena (Lair): {arena.name}
Tags: {', '.join(arena.tags) if arena.tags else '—'}

Villain A: {a.label()} — powers: {', '.join(a.powers) or '—'}
Villain B: {b.label()} — powers: {', '.join(b.powers) or '—'}

Return:
1) ARENA_SIGNATURE (3–6 bullets).
2) PROP_CATALOG (~12 items). Each: name, category, status_start, mass, interaction_hooks, uniqueness_rule, round_weight 1–3.
3) ENV_EVENTS (2–3 optional).
Constraints:
- Items must physically belong here; picked up/dragged/thrown if used.
- Keep it short and practical."""
    return [{"role":"system","content":sys}, {"role":"user","content":user}]

def _choose_props_for_round(ledger: dict, max_props: int = 4) -> list:
    catalog = ledger.get("PROP_CATALOG") or []
    used = ledger.setdefault("prop_uses", {})
    cooldowns = ledger.setdefault("prop_cooldowns", {})
    destroyed = set(ledger.get("destroyed_props", []))
    cand = []
    for item in catalog:
        name = (item.get("name") or "").strip()
        if not name or name in destroyed: continue
        if cooldowns.get(name, 0) > 0: continue
        try: weight = int(item.get("round_weight") or 1)
        except: weight = 1
        rule = (item.get("uniqueness_rule") or "").lower()
        times_used = used.get(name, 0)
        if "single" in rule and times_used >= 1: continue
        if "2 uses" in rule and times_used >= 2: continue
        cand.extend([item]*max(1,weight))
    random.shuffle(cand)
    picked, seen = [], set()
    for it in cand:
        if len(picked) >= max_props: break
        nm = it.get("name"); 
        if nm and nm not in seen:
            picked.append(it); seen.add(nm)
    ledger["props_in_play"] = [it.get("name") for it in picked if it.get("name")]
    return picked

# --------------- Parse labeled output ---------------
def _parse_round_text(txt: str) -> Tuple[str,str,Optional[str]]:
    a_text=b_text=""; camera=None
    t = txt.strip()
    cam = re.search(r'(^|\n)\s*CAMERA:\s*(.+)', t, flags=re.IGNORECASE)
    if cam:
        camera = "CAMERA: " + cam.group(2).strip()
        t = re.sub(r'(^|\n)\s*CAMERA:\s*.+(\n|$)', '\n', t, flags=re.IGNORECASE)
    a_match = re.search(r'\bA:\s*(.+?)(?:\n\s*B:|\Z)', t, flags=re.IGNORECASE|re.DOTALL)
    b_match = re.search(r'\bB:\s*(.+)$', t, flags=re.IGNORECASE|re.DOTALL)
    if a_match: a_text = a_match.group(1).strip()
    if b_match: b_text = b_match.group(1).strip()
    if not a_text or not b_text:
        paras = [p.strip() for p in re.split(r'\n\s*\n', t) if p.strip()]
        if len(paras)>=2: a_text=a_text or paras[0]; b_text=b_text or paras[1]
        elif len(paras)==1:
            s = re.split(r'(?<=[.!?])\s+', paras[0]); mid=max(1,len(s)//2)
            a_text=a_text or " ".join(s[:mid]); b_text=b_text or " ".join(s[mid:])
    return a_text.strip(), b_text.strip(), (camera.strip() if camera else None)

# --------------- Continuity auto-parser ---------------
INJURY_PATTERNS = [
    (r'\b(rib|ribs).*crack', 'cracked ribs'),
    (r'\bjaw\b.*(break|shatter|crack)', 'broken jaw'),
    (r'\b(concuss|dazed|blacked out)', 'concussion symptoms'),
    (r'\b(stab|impal|skewer)', 'stab wound'),
    (r'\b(bleed|blood|gush|spray)', 'bleeding'),
    (r'\b(laceration|gash|slash)', 'deep laceration'),
    (r'\b(dislocat|sprain)', 'dislocation/sprain'),
]
KNOCKDOWN_PATTERNS = [
    r'\b(prone|sprawl|face[- ]?down|on the ground|knocks? .* down)\b',
]
PROP_PICK_PAT = r'(grabs|snatches|rips|yanks|picks up|wields)\s+(the\s+)?(?P<prop>[\w\s\-]+)'
PROP_BREAK_PAT = r'(?P<prop>[\w\s\-]+)\s+(shatter|break|snap|splinter)s?'

HAZARD_SPAWNERS = [
    (r'\bglass\b', 'glass shards on floor'),
    (r'\b(wire|cable).*(sparks?)', 'sparking cables'),
    (r'\b(slick|slime|oil|blood)\b', 'slick ground'),
    (r'\bacid|toxin|chemical\b', 'corrosive puddle'),
]

def _append_unique(lst: List[str], item: str, max_len: int = 10):
    if item and item not in lst:
        lst.append(item)
        if len(lst) > max_len:
            lst.pop(0)

def _update_continuity_from_panels(a_text: str, b_text: str, ledger: dict, a_label: str, b_label: str):
    text_map = [(a_label, a_text), (b_label, b_text)]
    injuries = ledger.setdefault("injuries", {a_label:[], b_label:[]})
    hazards = ledger.setdefault("hazards", [])
    props_in_play = ledger.get("props_in_play", [])
    destroyed = ledger.setdefault("destroyed_props", [])
    in_hand = ledger.setdefault("in_hand", {a_label:None, b_label:None})
    prop_uses = ledger.setdefault("prop_uses", {})
    openings = ledger.setdefault("openings", {a_label:"", b_label:""})

    # Injuries + knockdowns
    for label, txt in text_map:
        lt = txt.lower()

        # injuries
        for pat, tag in INJURY_PATTERNS:
            if re.search(pat, lt):
                _append_unique(injuries.setdefault(label, []), tag)

        # knockdown -> positional disadvantage next round
        for pat in KNOCKDOWN_PATTERNS:
            if re.search(pat, lt):
                ledger["positional_disadvantage"] = f"{label} is prone and must recover before acting."
                break

        # openings
        if any(w in lt for w in ["stagger","reeling","dazed","winded","off-balance","opens","opening","limp","hobble","wobble"]):
            openings[label] = "off-balance; vulnerable next beat"

        # prop pick-ups
        pick = re.search(PROP_PICK_PAT, lt)
        if pick:
            # try to match to props_in_play name loosely
            cand = pick.group('prop').strip().lower()
            match = next((p for p in props_in_play if cand in p.lower()), None)
            if match:
                in_hand[label] = match
                prop_uses[match] = prop_uses.get(match,0) + 1

        # prop breaks
        br = re.search(PROP_BREAK_PAT, lt)
        if br:
            cand = br.group('prop').strip().lower()
            match = next((p for p in props_in_play if cand in p.lower()), None)
            if match and match not in destroyed:
                destroyed.append(match)

        # hazards spawned
        for pat, hz in HAZARD_SPAWNERS:
            if re.search(pat, lt):
                _append_unique(hazards, hz, max_len=8)

# --------------- Engine ---------------
def run_duel(a: Villain, b: Villain, arena: Arena, rounds: int = ROUNDS) -> DuelResult:
    if not USE_API:
        print("[DEBUG] DUEL_USE_OPENAI must be enabled.", file=sys.stderr)
        sys.exit(1)
    client = _client()

    ledger: Dict[str, any] = {
        "injuries": {a.label(): [], b.label(): []},
        "hazards": [],
        "props_in_play": [],
        "openings": {a.label(): "", b.label(): ""},
        "banlist": [],
        "positional_disadvantage": "",
        "destroyed_props": [],
        "in_hand": {a.label(): None, b.label(): None},
        "prop_uses": {},
        "prop_cooldowns": {},
    }

    # Scene
    scene = _chat_call(client, _scene_setter_messages(a,b,arena), max_tokens=450)

    # Prop Pack
    prop_text = _chat_call(client, _prop_pack_messages(a, b, arena), max_tokens=520)
    ledger["PROP_PACK_RAW"] = prop_text
    ledger["ARENA_SIGNATURE"] = []
    ledger["PROP_CATALOG"] = []
    ledger["ENV_EVENTS"] = []
    # Naive parse
    lines = [ln.strip("-• ").strip() for ln in prop_text.splitlines() if ln.strip()]
    bucket = None
    for ln in lines:
        lo = ln.lower()
        if "arena_signature" in lo: bucket="sig"; continue
        if "prop_catalog" in lo: bucket="cat"; continue
        if "env_events" in lo or "environment" in lo: bucket="ev"; continue
        if bucket=="sig": ledger["ARENA_SIGNATURE"].append(ln)
        elif bucket=="cat":
            item={"raw":ln}
            try:
                for p in [p.strip() for p in ln.split(",")]:
                    if ":" in p:
                        k,v = p.split(":",1)
                        item[k.strip().lower()] = v.strip()
                item["name"]=item.get("name") or item.get("raw")
            except: item["name"]=item.get("raw")
            ledger["PROP_CATALOG"].append(item)
        elif bucket=="ev": ledger["ENV_EVENTS"].append(ln)

    # Fight state
    a_total=b_total=0
    a_combo=b_combo=0
    a_stagger=b_stagger=0
    rng = random.Random()
    panels: List[RoundResult] = []

    TACTICS = ["feint & counter", "trap setup", "grapple/clinche", "mobility burst", "area denial", "terrain combo", "improvised weapon"]

    for r in range(1, rounds+1):
        # Position disadvantage tax
        if ledger.get("positional_disadvantage"):
            if a.label() in ledger["positional_disadvantage"]:
                a_stagger = max(a_stagger,1)
            if b.label() in ledger["positional_disadvantage"]:
                b_stagger = max(b_stagger,1)

        # Foreground props for this round
        _choose_props_for_round(ledger, max_props=4)

        # Random tactic bias per villain
        tactic_a = random.choice(TACTICS)
        tactic_b = random.choice(TACTICS)

        # CIC: compute extra injury penalties
        pen_a = _injury_penalty(a.label(), ledger, tactic_a)
        pen_b = _injury_penalty(b.label(), ledger, tactic_b)

        # Initiative
        a_first = (a_total >= b_total) if r > 2 else rng.random() < 0.5

        # Deltas (include CIC penalty)
        if a_first:
            a_delta,_ = _round_delta(a,b,arena,rng,a_stagger,a_combo, extra_penalty=pen_a)
            b_delta,_ = _round_delta(b,a,arena,rng,b_stagger,b_combo, extra_penalty=pen_b)
        else:
            b_delta,_ = _round_delta(b,a,arena,rng,b_stagger,b_combo, extra_penalty=pen_b)
            a_delta,_ = _round_delta(a,b,arena,rng,a_stagger,a_combo, extra_penalty=pen_a)

        # Totals & streaks
        a_total += a_delta; b_total += b_delta
        a_combo = (a_combo+1) if a_delta>0 else 0
        b_combo = (b_combo+1) if b_delta>0 else 0
        a_stagger = 1 if (b_delta - a_delta) >= 5 else 0
        b_stagger = 1 if (a_delta - b_delta) >= 5 else 0

        # Round call
        msgs = _round_messages(a,b,arena,ledger,r,a_total,b_total,rounds,tactic_a,tactic_b)
        txt = _chat_call(client, msgs, max_tokens=680)
        a_panel, b_panel, camera = _parse_round_text(txt)

        # Update continuity from narration
        _update_continuity_from_panels(a_panel, b_panel, ledger, a.label(), b.label())

        panels.append(RoundResult(
            r=r,
            a_text=_wrap(a_panel),
            b_text=_wrap(b_panel),
            camera=camera,
            a_delta=a_delta, b_delta=b_delta,
            a_total=a_total, b_total=b_total
        ))

        # Cooldown positional disadvantage after one round (fighter recovered)
        ledger["positional_disadvantage"] = ""

        # Simple prop cooldown tick
        for k in list(ledger.get("prop_cooldowns", {}).keys()):
            ledger["prop_cooldowns"][k] = max(0, ledger["prop_cooldowns"][k]-1)

    # Winner & finisher
    winner, loser = (a,b) if a_total>b_total else (b,a) if b_total>a_total else ((a,b) if panels[-1].a_delta>=panels[-1].b_delta else (b,a))
    finisher = _chat_call(client, _finisher_messages(winner, loser, arena, ledger), max_tokens=420)

    return DuelResult(a=a,b=b,arena=arena,scene_setter=scene,rounds=panels,
                      a_total=a_total,b_total=b_total,winner_label=winner.label(),
                      finisher=_wrap(finisher))

# --------------- Output ---------------
def print_duel(dr: DuelResult) -> None:
    print(f"=== {dr.a.label()} vs {dr.b.label()} — {dr.arena.name} ===\n")
    print(_wrap(dr.scene_setter)); print()
    for rr in dr.rounds:
        print(f"Round {rr.r}")
        if rr.camera:
            print(rr.camera if rr.camera.startswith("CAMERA:") else f"CAMERA: {rr.camera}")
        print(rr.a_text)
        print(rr.b_text)
        print(f"- Score Change: {_name(dr.a)} {rr.a_delta:+}, {_name(dr.b)} {rr.b_delta:+}")
        print(f"- Totals: {_name(dr.a)} {rr.a_total} — {_name(dr.b)} {rr.b_total}")
        print()
    print(f"Winner: {dr.winner_label}")
    print("Finisher:")
    print(dr.finisher)

# --------------- DOCX Export ---------------
def export_duel_to_docx(dr: DuelResult, out_dir: str = DUEL_DOCX_DIR) -> Optional[str]:
    try:
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except Exception as e:
        print(f"[WARN] python-docx not installed; skip .docx export. ({e})")
        return None

    # Resolve and create directory
    try:
        from pathlib import Path
        Path(out_dir).mkdir(parents=True, exist_ok=True)
    except Exception as e:
        print(f"[WARN] Could not create output directory: {out_dir} ({e})")
        return None

    # Build document
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # Title
    p = doc.add_paragraph()
    r = p.add_run(f"{dr.a.label()} vs {dr.b.label()} — {dr.arena.name}")
    r.bold = True; r.font.size = Pt(22)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # Scene Intro
    h = doc.add_paragraph(); hr = h.add_run("Scene Intro"); hr.bold = True; hr.font.size = Pt(14)
    doc.add_paragraph(dr.scene_setter); doc.add_paragraph()

    # Rounds
    for rr in dr.rounds:
        doc.add_paragraph().add_run("").add_break()
        h = doc.add_paragraph(); t = h.add_run(f"Round {rr.r}"); t.bold=True; t.font.size = Pt(14)
        if rr.camera:
            pcam = doc.add_paragraph(); rcam = pcam.add_run(rr.camera); rcam.italic = True
        # A block
        pa = doc.add_paragraph(); ra = pa.add_run(dr.a.label()); ra.bold = True
        doc.add_paragraph(rr.a_text)
        # B block
        pb = doc.add_paragraph(); rb = pb.add_run(dr.b.label()); rb.bold = True
        doc.add_paragraph(rr.b_text)
        # Scores
        sc = doc.add_paragraph(); r1 = sc.add_run("- Score Change: "); r1.bold=True
        sc.add_run(f"{_name(dr.a)} {rr.a_delta:+}, {_name(dr.b)} {rr.b_delta:+}")
        st = doc.add_paragraph(); r2 = st.add_run("- Totals: "); r2.bold=True
        st.add_run(f"{_name(dr.a)} {rr.a_total} — {_name(dr.b)} {rr.b_total}")

    # Winner / Finisher
    doc.add_paragraph().add_run("").add_break()
    hw = doc.add_paragraph(); rw = hw.add_run("Winner"); rw.bold = True; rw.font.size = Pt(14)
    w = doc.add_paragraph(); wr = w.add_run(dr.winner_label); wr.bold = True
    doc.add_paragraph().add_run("").add_break()
    hf = doc.add_paragraph(); rf = hf.add_run("Finisher"); rf.bold = True; rf.font.size = Pt(14)
    doc.add_paragraph(dr.finisher)

    # Filename
    ts = _dt.datetime.now().strftime("%Y-%m-%d_%H%M")
    safe_arena = re.sub(r'[^A-Za-z0-9 _-]+', '', dr.arena.name).strip().replace(" ", "_")
    safe_a = re.sub(r'[^A-Za-z0-9 _-]+', '', dr.a.label()).strip().replace(" ", "_")
    safe_b = re.sub(r'[^A-Za-z0-9 _-]+', '', dr.b.label()).strip().replace(" ", "_")
    filename = f"{ts}_{safe_a}_vs_{safe_b}_{safe_arena}.docx"

    out_path = os.path.join(out_dir, filename)
    try:
        doc.save(out_path)
        print(f"[OK] Saved duel story: {out_path}")
        return out_path
    except Exception as e:
        print(f"[WARN] Failed to save .docx: {e}")
        return None

# --------------- Loading ---------------
def load_villain_from_json(path: str) -> Tuple[Villain, str]:
    if not os.path.exists(path):
        raise FileNotFoundError(f"[DEBUG] File not found: {path}")
    with open(path, "r", encoding="utf-8") as f:
        data = json.load(f)
    name  = data.get("name") or "[Unnamed]"
    alias = data.get("alias") or None
    powers = data.get("power") or data.get("powers") or ""
    powers = [powers] if isinstance(powers, str) else (powers or [])
    weaknesses = data.get("weakness") or data.get("weaknesses") or ""
    weaknesses = [weaknesses] if isinstance(weaknesses, str) else (weaknesses or [])
    catchphrase = data.get("catchphrase") or None
    gender = data.get("gender") or ""
    subj, obj, pos = _pronouns_from_gender(gender)
    v = Villain(name=name, alias=alias, powers=powers, weaknesses=weaknesses,
                catchphrase=catchphrase, vibe=data.get("theme") or None,
                subj=subj, obj=obj, pos=pos)
    return v, (data.get("lair") or "").strip()

def default_villains() -> Tuple[Villain, Villain, Arena]:
    parser = argparse.ArgumentParser(add_help=False)
    parser.add_argument("--a", type=str); parser.add_argument("--b", type=str)
    args, _ = parser.parse_known_args()
    if args.a and args.b:
        va, lair_a = load_villain_from_json(args.a)
        vb, lair_b = load_villain_from_json(args.b)
        chosen_lair = random.choice([lair_a, lair_b]) or "Unknown Arena"
        tags = KNOWN_LAIR_TAGS.get(chosen_lair) or _auto_tags(chosen_lair)
        return va, vb, Arena(name=chosen_lair, tags=tags)
    # Demo fallback
    aria = Villain(name="Aria Greene", alias="Mimic Mistress",
        powers=["Shapeshifting at cellular level; can grow bone/talon weapons."],
        weaknesses=["Extreme cold slows regeneration."],
        catchphrase="Faces shift, truths blur", vibe="deceiver",
        subj="She", obj="her", pos="her")
    chem = Villain(name="Benjamin Silva", alias="The Chem Burner",
        powers=["Green toxin from hands; dissolves flesh, bone, and metal."],
        weaknesses=["Neutralized by specialized antitoxin rigs."],
        catchphrase="Feel the burn of my touch", vibe="industrial",
        subj="He", obj="him", pos="his")
    lairs = [
        {"owner":"Aria","name":"BioLab Sanctum","tags":["lab","biotech","low_light","glass","props","sterile"]},
        {"owner":"Chem","name":"Acid Den Hideout","tags":["industrial","metal","steam","enclosed","props","toxic"]},
    ]
    chosen = random.choice(lairs)
    return aria, chem, Arena(name=chosen["name"], tags=chosen["tags"])

# --------------- Main ---------------
def main():
    a, b, arena = default_villains()
    dr = run_duel(a,b,arena, rounds=ROUNDS)
    print_duel(dr)
    # Auto-export to DOCX after printing
    export_duel_to_docx(dr, DUEL_DOCX_DIR)

if __name__ == "__main__":
    main()
